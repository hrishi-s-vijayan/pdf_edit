import subprocess
import tabula

from django.http import HttpResponse
from pypdf import PdfWriter, PdfReader

from docx import Document
from docx.shared import Pt, Inches

from pathlib import Path
from pdf2docx import Converter

from pypdf import PdfWriter

from .utils import find_font_size_of_text


from spire.pdf.common import *
from spire.pdf import *


import aspose.pdf as ap



import fitz

import pdfrw


pdf_writer = PdfWriter()

pdf_path = (
    Path.home()
    / "Desktop/Hrishi/Django_projects/django_code_optimization/pdf_app/anirudh.pdf"
    # / "practice_files"
    # / "Pride_and_Prejudice.pdf"
)

def generate_pdf(request):
    # # Create a new PDF
    # pdf = PdfWriter()

    # # Add a page
    # pdf.addPage()

    # # Set content
    # pdf.setFont("Arial", 12)
    # pdf.drawString(100, 750, "Hello, World!")

    # # Save the PDF to a buffer
    # pdf_buffer = HttpResponse(content_type='application/pdf')
    # pdf.write(pdf_buffer)

    # # Create a response
    # pdf_buffer['Content-Disposition'] = 'attachment; filename="example.pdf"'
    # return pdf_buffer

    #########################################################################

    # pdf_reader = PdfReader(pdf_path)
    # length = len(pdf_reader.pages)

    # print('========== length is : ', length)

    # meta_data = pdf_reader.metadata

    # print('========= meta data is : ', meta_data)

    # title = pdf_reader.metadata.title
    
    # print('========== title is : ', title)

    # first_page = pdf_reader.pages[0]
    # page_type = type(first_page)

    # print('========= first page is : ', first_page)

    # print('=========== page type is : ', page_type)

    # print("======= text is : ", first_page.extract_text())

    # # for page in pdf_reader.pages:
    # #     print("======= ", page.extract_text())


    # txt_file = Path.home() / "Desktop/Hrishi/Django_projects/django_code_optimization/pdf_app/test.txt"
    # content = [
    #     f"{pdf_reader.metadata.title}",
    #     f"Number of pages: {len(pdf_reader.pages)}"
    # ]

    # for page in pdf_reader.pages:
    #     content.append(page.extract_text())

    # txt_file.write_text("\n".join(content))

    # page = pdf_writer.add_blank_page(width=8.27 * 72, height=11.7 * 72)

    # type(page)

    # # pdf_writer.write("blank2.pdf")

    # input_pdf = PdfReader(pdf_path)

    # first_page = input_pdf.pages[0]

    # print('===== first page is : ', first_page)

    # output_pdf = PdfWriter()
    # output_pdf.add_page(first_page)

    # output_pdf.write("first_page.pdf")

    ###################################################################

    ############**************######################################################################

    # Create an object of the PdfDocument class
    # doc = PdfDocument()
    # # # Load a PDF file
    # # doc.LoadFromFile("first_page.pdf")
    # doc.LoadFromFile("test_doc.pdf")

    # # Iterate through the pages in the document
    # for i in range(doc.Pages.Count):
    #     # Get the current page
    #     page = doc.Pages[i]    
    #     # Create an object of the PdfTextReplace class and pass the page to the constructor of the class as a parameter
    #     replacer = PdfTextReplacer(page)
        
    #     # Replace All instances of a specific text with new text
    #     # replacer.ReplaceAllText("Signedly", "replaced text tttttttttt")
    #     replacer.ReplaceAllText("newsletter", "TTTTTT")

    # #     # Replace All instances of a specific text with new text and set text color
    # #     # replacer.ReplaceAllText("PDF Editor", "PDF Editor", Color.get_Yellow())

    # # Save the resulting file
    # doc.SaveToFile("ReplaceAllFoundText.pdf")

    # doc.Close()


    ############**************######################################################################
    

    ###############-----------------------------------------###############################################

    # text_to_add = request.POST.get('textToAdd', 'hellooooooooooo')
    # x_pos = float(request.POST.get('xPos', 7))
    # y_pos = float(request.POST.get('yPos', 20))
    # page_number = request.POST.get('page_number', 0)

    # pdf_path = "first_page.pdf"

    # output_path = "modified_example.pdf"

    # # Edit PDF using PyMuPDF
    # pdf_document = fitz.open(pdf_path)
    # page = pdf_document[page_number]  # For simplicity, assuming you're editing the first page

    # text_instances = page.search_for("Signedly")
    
    # # Print text and its coordinates
    # for text, rect in zip(page.get_text().splitlines(), text_instances):
    #     print(f"Text: {text}, Position: {rect}")
    

    
    # # Add text to the specified position
    # # page.insert_text((x_pos, y_pos), text_to_add)

    # text_box = fitz.Rect(x_pos, y_pos, x_pos + 200, y_pos + 20)  # Adjust dimensions as needed
    # text_color = (0, 0, 0)

    # # Set font parameters
    # # font_name = "Arial"  # Example font name
    # font_name = "Helvetica"
    # font_size = 120  # Example font size

    # page.insert_text(((7.0, 20.093)), text_to_add)

    # # page.insert_textbox(text_box, text_to_add, fontsize=font_size, fontname=font_name, color=text_color)


    # pdf_document.save(output_path)
    # pdf_document.close()

    ###############--------------------------------###############################################

    ##################@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@############################################

    # input_pdf_path = "first_page.pdf"
    # output_pdf_path = "new.pdf"

    # pdf_document = fitz.open(input_pdf_path)
    
    # # Create a new PDF document
    # new_pdf = fitz.open()
    
    # # Iterate over pages in the existing PDF
    # for page_number in range(pdf_document.page_count):
    #     # Get the page
    #     page = pdf_document[page_number]
        
    #     # Create a new page in the new PDF and insert the existing page
    #     new_page = new_pdf.new_page(width=page.rect.width, height=page.rect.height)
    #     new_page.insert_page(page_number, page)
    
    # # Save the new PDF document
    # new_pdf.save(output_pdf_path)
    
    # # Close the PDFs
    # pdf_document.close()
    # new_pdf.close()

    ##################@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@############################################

    import PyPDF2

    input_pdf_path = "first_page.pdf"
    output_pdf_path = "final_output.pdf"
    search_text = "Signedly"
    replace_text = "yldengiS"

    ###############################################################

    # # Create a PDF writer object
    # pdf_writer = PyPDF2.PdfWriter()

    # with open(input_pdf_path, "rb") as input_file:
    #     # Create a PDF reader object
    #     pdf_reader = PyPDF2.PdfReader(input_file)
        
        
        
    #     # Iterate over each page in the PDF
    #     for page_number in range(len(pdf_reader.pages)):
    #         # Get the current page
    #         page = pdf_reader.pages[page_number]
            
    #         # Search for the search text and replace it with the replace text
    #         text = page.extract_text()
    #         if search_text in text:
    #             text = text.replace(search_text, replace_text)
    #             # page.mergePage(PyPDF2.PageObject.create_text_object(text))
    #             # Create a new page object
    #             new_page = PyPDF2.PageObject.create_blank_page(width=page.mediaBox.getWidth(), height=page.mediaBox.getHeight())
            
    #         # Set the text content of the new page
    #         new_page.mergePage(page)
    #         new_page.addText(new_text_content)
    
            
    #         # Add the modified page to the PDF writer
    #         pdf_writer.add_page(page)
        
    #     # Write the modified PDF to the output file
    #     with open(output_pdf_path, "wb") as output_file:
    #         pdf_writer.write(output_file)

    ##########################################################################
    # **************************************************************************

    # it works but won't work with pdfs generated from pdf
    # from .solution5 import search_and_replace_text
    
    # search_and_replace_text(search_text, replace_text)
    
    # **************************************************************************
    ###########################################################################


    ############################################################################
    #222222222222222222222222222222222222222222222222222222222222222222222222222

    # input_pdf_path = "test_doc.pdf"
    # # input_pdf_path = "first_page.pdf"
    # # output_pdf_path = "output_redacted.pdf"
    # search_text = "Your goal is to showcase"
    # replace_text = "showcase to is your goal"

    # from .pymupdf_solution2 import replace_text_with_matching_font

    # replace_text_with_matching_font(input_pdf_path, output_pdf_path, search_text, replace_text)

    #222222222222222222222222222222222222222222222222222222222222222222222222222
    ############################################################################


    ############################################################################
    #333333333333333333333333333333333333333333333333333333333333333333333333333

    from pdf2docx import Converter
    from docx.shared import Pt

    # import aspose.words as aw

    

    def convert_pdf_to_word(input_pdf_path, output_docx_path):
        cv = Converter(input_pdf_path)
        cv.convert(output_docx_path)
        cv.close()

    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # def adjust_table_alignment(input_docx_path, output_docx_path):
    #     """
    #     Adjusts the alignment of tables in a DOCX file to make them closer to the left margin.

    #     Parameters:
    #         input_docx_path (str): The path to the input DOCX file.
    #         output_docx_path (str): The path to save the adjusted DOCX file.
    #     """
    #     # Load the DOCX file
    #     doc = Document(input_docx_path)

    #     # Iterate through tables in the document
    #     for table in doc.tables:
    #         print('======== table is : ', table)
    #         # Set the table alignment to left
    #         table.alignment = WD_TABLE_ALIGNMENT.CENTER

    #     # Save the modified DOCX file
    #     doc.save(output_docx_path)


    def adjust_table_alignment(input_docx_path, output_docx_path, content_alignment='left'):
        """
        Adjusts the alignment of tables and their contents in a DOCX file.

        Parameters:
            input_docx_path (str): The path to the input DOCX file.
            output_docx_path (str): The path to save the adjusted DOCX file.
            content_alignment (str): Desired content alignment. Options: 'left', 'center', 'right'.
                                    Defaults to 'left'.
        """
        # Load the DOCX file
        doc = Document(input_docx_path)

        # Set paragraph alignment based on user input
        if content_alignment == 'left':
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif content_alignment == 'center':
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif content_alignment == 'right':
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid content alignment. Use 'left', 'center', or 'right'.")

        # Iterate through tables in the document
        for table in doc.tables:
            
            
            table = remove_whitespaces_from_table(table)
            widht_list = max_text_length_per_column(table)
            table = adjust_table_column_width(table, widht_list)
            

            print("WWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWW width list is : ", widht_list)
            
            

            # print('@@@@@@@@@@@@@@@@@@@@@@@@@@', table.rows[0].cells)

            # if len(table.rows[0].cells) == 2:
            #     # print('############# inside it ')
            #     # column_widths = [1000000, 500000]
            #     # adjust_column_widths(table, column_widths)
            #     # set_table_layout(table, 'auto') 
            #     doc.save(output_docx_path)

            # Set the table alignment to center
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            

            max_cell_info, max_length, max_cell_text = find_cell_with_maximum_length(table)

            print("..........  max lenght is : ", max_length, max_cell_text)
            
            # Iterate through each row in the table
            for row in table.rows:
                # Iterate through each cell in the row
                for cell in row.cells:
                    print("======= cell is : ", cell.text, 'lenght is : ', len(cell.text), 'type is : ', type(cell.text))

                    # cell.text = cell.text.lstrip()

                    # if len(cell.text) > 0:
                    #     print("======= first letter is : ", cell.text[0])

                    if (len(cell.text) < max_length ):
                        # print('kkkkkkkk test cell is : ', cell.text)
                        cell.text = cell.text + generate_trailing_empty_space_string(max_length - len(cell.text))
                    
                    # if len(cell.text) == 1:
                        
                    #     cell.text = generate_empty_space_string(max_length) + cell.text
                    #     print('lenght is : ', len(generate_empty_space_string(max_length)))


                    # if len(cell.text) < max_length:
                    #     cell.text = cell.text + generate_empty_space_string(max_length - len(cell.text))

                    # print("%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%% grid span is : ", grid_span)
                    # Align the content in each cell
                    for paragraph in cell.paragraphs:
                        # print('====== cell paragraph is : ', paragraph.alignment, type(paragraph.alignment), paragraph.paragraph_format.right_indent)
                        # print('.......... alignemnt type is : ', type(paragraph.paragraph_format.right_indent))
                        # print('====== cell paragraph is : ', dir(paragraph))

                        # print('--------- space before is : ', paragraph.paragraph_format.space_before)
                        # print('--------------------- space after is : ', paragraph.paragraph_format.space_after)
                        # print('--------------------- line space is : ', paragraph.paragraph_format.line_spacing)

                        # print('00000000000000000 current alignment is : ', paragraph.alignment)

                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        paragraph.paragraph_format.right_indent = Pt(10)
                        # paragraph.paragraph_format.space_before = Pt(0)
                        # paragraph.paragraph_format.tab_stops.clear_all()

                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

        # Save the modified DOCX file
        doc.save(output_docx_path)

        #############################################################

        



    def pdf_to_rtf(input_pdf_path, output_rtf_path):
        # Create a PDF reader
        pdf_reader = PdfReader(input_pdf_path)
        
        # Initialize an empty list to store text content from each page
        text_content = []
        
        # Iterate through each page in the PDF
        for page in pdf_reader.pages:
            # Extract text from the page
            text = page.extract_text()
            text_content.append(text)
        
        # Write the extracted text content to an RTF file
        with open(output_rtf_path, 'w') as rtf_file:
            # Write the RTF header
            rtf_file.write("{\\rtf1\\ansi\\deff0\n")
            
            # Write the text content
            for text in text_content:
                # RTF special characters must be escaped
                text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
                rtf_file.write(text)
            
            # Write the RTF footer
            rtf_file.write("\n}")
        
        print(f"PDF content has been converted to RTF and saved to {output_rtf_path}")

    
    def rtf_to_docx_libreoffice(input_rtf_path, output_docx_path):
        """Converts an RTF file to a DOCX file using LibreOffice.

        Args:
            input_rtf_path (str): Path to the input RTF file.
            output_docx_path (str): Path to save the output DOCX file.
        """
        print('======= inside libre office : ')
        # command = f"libreoffice --headless --convert-to docx:writer_microsoft_2007 {input_rtf_path} --outdir {output_docx_path}"
        # subprocess.run(command.split(), shell=False, check=True)
        try:
            # command = f"libreoffice --headless --convert-to docx:writer_microsoft_2007 {input_rtf_path} --outdir {output_docx_path}"
            command = ["libreoffice", "--headless", "--convert-to", "docx:writer_microsoft_2007", input_rtf_path, "--outdir", output_docx_path]
            subprocess.run(command, shell=False, check=True)
            print(f"RTF file converted to DOCX using LibreOffice: {output_docx_path}")
        except subprocess.CalledProcessError as e:
            print(f"Error: LibreOffice conversion failed. Error code: {e.returncode}")

        # command = ["libreoffice", "--headless", "--convert-to", "docx:writer_microsoft_2007", input_rtf_path, "--outdir", os.path.dirname(output_docx_path)]




    # input_pdf_path = "test_doc.pdf"
    input_pdf_path = "table.pdf"
    # input_pdf_path = "Hrishi_Doc.pdf"
    # input_pdf_path = "ebook1.pdf"
    # input_pdf_path = "sample2.pdf"
    output_docx_path = "output.docx"

    convert_pdf_to_word(input_pdf_path, output_docx_path)

    # change_page_size("output.docx", "output.docx", page_width=8.5, page_height=11)
    adjust_margins(output_docx_path)

    reduce_font_size(output_docx_path)

    adjust_table_alignment(output_docx_path, output_docx_path)

    document_path = "output.docx"
    table_row_index = 0  # Zero-based index of the row containing the table
    column_index = 1     # Zero-based index of the column to adjust
    new_width_in_twips = 5000  # New width for the column in twips (adjust as needed)

    print('QQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQ  before')
    # adjust_column_width_openxml(document_path, table_row_index, column_index, new_width_in_twips)
    print('QQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQQ  after')

    # output_rtf_path = "output99.rtf"

    # pdf_to_rtf(input_pdf_path, output_rtf_path)

    # rtf_to_docx_libreoffice("output.rtf", output_docx_path)

    

    def convert_docx_to_pdf(input_docx_path, output_pdf_path):
        """
        Convert a DOCX file to a PDF file using LibreOffice.

        Parameters:
        input_docx_path (str): The path to the input DOCX file.
        output_pdf_path (str): The path where the output PDF file will be saved.
        """
        # Use subprocess to call the libreoffice command for conversion
        command = ["libreoffice", "--headless", "--convert-to", "pdf", input_docx_path, "--outdir", os.path.dirname(output_pdf_path)]
        
        try:
            subprocess.run(command, check=True)
            print(f"Successfully converted {input_docx_path} to {output_pdf_path}")
        except subprocess.CalledProcessError as e:
            print(f"Failed to convert {input_docx_path} to {output_pdf_path}: {e}")



    # def edit_word_document(input_docx_path, output_docx_path, search_text, replace_text, font_size, is_bold=False):
    #     doc = Document(input_docx_path)

    #     # Iterate through each paragraph in the document
    #     for paragraph in doc.paragraphs:
    #         # Create a list to hold new runs
    #         new_runs = []

    #         # Iterate through the runs in the paragraph
    #         for run in paragraph.runs:
    #             # Get the text of the run
    #             text = run.text
    #             # Start index for searching the word to bold
    #             start_index = 0

    #             # Loop through occurrences of the word to replace
    #             while True:
    #                 # Find the word in the current run text starting from the start index
    #                 start_index = text.find(search_text, start_index)

    #                 # If the word is not found, append the remaining text as a single run and break the loop
    #                 if start_index == -1:
    #                     # Append the remaining text as a new run with the current run's formatting
    #                     if text:
    #                         new_runs.append((text, run.bold, run.font.size))
    #                         pass
    #                     break

    #                 # Append the text before the word as a new run with the current run's formatting
    #                 before_word = text[:start_index]
    #                 if before_word:
    #                     new_runs.append((before_word, run.bold, run.font.size))

    #                 # Append the replacement word as a new run with specified boldness and font size
    #                 replaced_word = replace_text
    #                 new_runs.append((replaced_word, is_bold, Pt(font_size)))

    #                 # Update start index and remaining text
    #                 start_index += len(search_text)
    #                 text = text[start_index:]

    #         # Clear the existing runs in the paragraph
    #         paragraph.clear()

    #         # Add the new runs to the paragraph with the correct formatting
    #         for run_text, bold, size in new_runs:
    #             new_run = paragraph.add_run(run_text)
    #             new_run.bold = bold
    #             new_run.font.size = size

    #     # Save the modified document
    #     doc.save(output_docx_path)

    ######## NEW #########
    def edit_word_document(input_docx_path, output_docx_path, search_text, replace_text, font_size, is_bold=False):
        doc = Document(input_docx_path)

        # Iterate through each paragraph in the document
        for paragraph in doc.paragraphs:
            # Iterate through the runs in the paragraph
            for run in paragraph.runs:
                # Get the text of the run
                text = run.text

                # Find the word in the current run text
                start_index = text.find(search_text)
                if start_index != -1:  # If the word is found
                    # Replace the word with the new text
                    replaced_text = text[:start_index] + replace_text + text[start_index + len(search_text):]

                    # Clear the existing run and add the replaced text
                    run.clear()
                    run.text = replaced_text

                    # Apply formatting to the replaced text
                    run.bold = is_bold
                    run.font.size = Pt(font_size)

        # Save the modified document
        doc.save(output_docx_path)
    

    ######### NEW 2 ###########
    # def edit_word_document(input_docx_path, output_docx_path, search_text, replace_text, font_size, is_bold=False):
    #     doc = Document(input_docx_path)

    #     # Iterate through each paragraph in the document
    #     for paragraph in doc.paragraphs:
    #         # Iterate through the runs in the paragraph
    #         for run in paragraph.runs:
    #             # Get the text of the run
    #             text = run.text

    #             # Find the word in the current run text
    #             start_index = text.find(search_text)
    #             if start_index != -1:  # If the word is found
    #                 # Split the run into three parts: before, search text, and after
    #                 before_text = text[:start_index]
    #                 search_text_part = text[start_index:start_index + len(search_text)]
    #                 after_text = text[start_index + len(search_text):]

    #                 # Clear the existing run
    #                 run.clear()

    #                 # Add the before text
    #                 run.add_text(before_text)

    #                 # Add the replaced text with formatting
    #                 replaced_run = run.add_run(replace_text)
    #                 replaced_run.font.size = Pt(font_size)
    #                 if is_bold:
    #                     replaced_run.bold = True

    #                 # Add the after text
    #                 run.add_text(after_text)

    #     # Save the modified document
    #     doc.save(output_docx_path)



    input_docx_path = "output.docx"
    output_docx_path = "edited_output.docx"
    search_text = "following data were collected"
    replace_text = "YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH YASH"
    # search_text = "asdf"
    # replace_text = "asdf"
    
    try:
        font_size, is_bold =  find_font_size_of_text(input_pdf_path, search_text)
    except:
        font_size, is_bold = 10, False

    print('======= TESTING WHETHER BOLD OR NOT : ', is_bold)

    is_word_exist = search_text_in_document(input_docx_path, search_text)

    print('........ is word exist : ', is_word_exist)

    if is_word_exist:
        edit_word_document(input_docx_path, output_docx_path, search_text, replace_text, font_size, is_bold)
        input_docx_path = "edited_output.docx"
    else:
        input_docx_path = "output.docx"

    # Usage
    # input_docx_path = "output.docx"
    # input_docx_path = "edited_output.docx"
    output_pdf_path = "final_output.pdf"

    convert_docx_to_pdf(input_docx_path, output_pdf_path)


    print('..............test size is : ', font_size)
    

    #333333333333333333333333333333333333333333333333333333333333333333333333333
    ############################################################################
    

    return HttpResponse('hh')

def find_cell_with_maximum_length(table):
    """
    Finds the cell with the maximum text length in the given table, skipping the first row.

    Parameters:
        table: The table object (from the `docx` library) to search.

    Returns:
        tuple: A tuple containing the row index and cell index of the cell with the maximum length.
               Also returns the maximum length and the text of the cell.
    """
    # Initialize variables to track the cell with the maximum length
    max_length = 0
    max_cell_info = (None, None)  # (row_index, cell_index)
    max_cell_text = ""
    
    # Iterate through each row in the table, starting from the second row (index 1)
    for row_index in range(1, len(table.rows)):
        row = table.rows[row_index]
        # Iterate through each cell in the row
        for cell_index, cell in enumerate(row.cells):
            # Calculate the length of the cell text
            cell_text_length = len(cell.text)
            
            # Check if this cell has the maximum length so far
            if cell_text_length > max_length:
                max_length = cell_text_length
                max_cell_info = (row_index, cell_index)
                max_cell_text = cell.text
    
    return max_cell_info, max_length, max_cell_text


def generate_empty_space_string(length: int) -> str:
    """
    Returns a string of empty spaces with the specified length.

    Parameters:
        length (int): The length of the string of empty spaces.

    Returns:
        str: A string of empty spaces with the specified length.
    """
    # Validate the input length
    if length < 0:
        raise ValueError("Length must be a non-negative integer.")
    
    # Return a string of empty spaces with the specified length
    return ' ' * length

def generate_trailing_empty_space_string(length: int) -> str:
    """
    Returns a string of empty spaces with the specified length.

    Parameters:
        length (int): The length of the string of empty spaces.

    Returns:
        str: A string of empty spaces with the specified length.
    """
    # Validate the input length
    if length < 0:
        raise ValueError("Length must be a non-negative integer.")
    
    # Return a string of empty spaces with the specified length
    return '\u2002' * length


def remove_whitespaces_from_table(table):
    """
    Removes all leading white spaces (spaces, tabs, newlines, etc.) from each cell's text in the given table.

    Parameters:
        table (docx.table.Table): The input table from which leading white spaces in cell text will be removed.

    Returns:
        docx.table.Table: The table with leading white spaces removed from each cell's text.
    """
    # Iterate through each row in the table
    for row in table.rows:
        # Iterate through each cell in the row
        for cell in row.cells:
            # Remove leading white spaces from the cell's text
            cell.text = cell.text.strip()
            
    # Return the modified table
    return table





from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_layout(table, layout_type):
    """
    Set the layout type of a table to either 'fixed' or 'auto'.

    Parameters:
        table (docx.table.Table): The table whose layout needs to be adjusted.
        layout_type (str): The desired layout type, either 'fixed' or 'auto'.
    """
    # Access the underlying XML element of the table
    table_element = table._element
    
    # Access or create the table properties (tblPr)
    if not hasattr(table_element, 'tblPr'):
        table_properties = OxmlElement('w:tblPr')
        table_element.append(table_properties)
    else:
        table_properties = table_element.tblPr
    
    # Create an OxmlElement for the table layout
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), layout_type)
    
    # Add the layout element to the table properties
    table_properties.append(tbl_layout)

    table.style = 'Colorful List'

    print('&&&&&&&&&&&&&&&&&&&&&&&&&&&&& table properties are : ', dir(table))


from docxtpl import DocxTemplate

def adjust_column_width_openxml(document_path, table_row_index, column_index, new_width_in_twips):
    """Adjusts the width of a specific column in a table within a DOCX document using OpenXML.

    Args:
        document_path (str): Path to the DOCX document.
        table_row_index (int): The zero-based index of the row containing the table.
        column_index (int): The zero-based index of the column to adjust.
        new_width_in_twips (int): The new width for the column in twips (1 twip = 1/20th of a point).
    """
    try:
        template = DocxTemplate(document_path)
        context = template.context

        # Access table using its name in the template (if defined)
        # Or, iterate through all tables to find the desired one
        table = context['table_name']  # Replace with actual table name in your template

        print('*************** table is : ', table)

        # Get table grid object (assumes a single table section)
        grid = table.xpath_find("w:tblGrid", first=True)

        # Access grid child element corresponding to the target column
        grid_col = grid.xpath_find(f"w:gridCol[@w:w={column_index+1}]", first=True)

        # Set the new width in twips
        grid_col.set('w:w', str(new_width_in_twips))

        # Render the modified context back to the document
        template.render(context)
        template.save(document_path)
        print(f"Column width adjusted for column {column_index+1} in table at row {table_row_index+1}.")
    except Exception as e:
        print(f"Error: {e}")

# def adjust_table_column_width(table):
#     # Iterate through each column in the table
#     for col_index in range(len(table.columns)):
#         # Get the current width of the column
#         current_width = table.columns[col_index].width

#         print('llllllllllllllllllllllllllllll current widht is : ', current_width, type(current_width))
        
#         # Double the width of the column
#         # new_width = int(current_width * 1)
#         new_width = int(132842)
#         table.columns[col_index].width = new_width
        
#         # Set the new width for each cell in the column
#         for cell in table.columns[col_index].cells:
#             print('mmmmmmmmmmmmmmmm cell text is : ', cell.text)
#             tc = cell._element
#             tcPr = tc.get_or_add_tcPr()
#             tcW = OxmlElement('w:tcW')
#             tcW.set(qn('w:w'), str(new_width))
#             tcW.set(qn('w:type'), 'dxa')
#             tcPr.append(tcW)

#     return table

def adjust_table_column_width(table, max_text_lengths):
    # Calculate the total maximum text length across all columns
    total_length = sum(max_text_lengths)
    
    # Iterate through each column in the table
    for col_index in range(len(table.columns)):
        # Calculate the proportion of the column width
        proportion = max_text_lengths[col_index] / total_length

        print('======== PROPORTION IS : ', proportion)
        
        # Set a base width, which you can adjust based on your requirements
        base_width = 7084200
        
        # Calculate the new width for the column based on the proportion
        new_width = int(base_width * proportion)
        
        # Update the column width
        table.columns[col_index].width = new_width
        
        # Set the new width for each cell in the column
        for cell in table.columns[col_index].cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(new_width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    return table


def max_text_length_per_column(table):
    # Initialize a list to store the maximum text length for each column
    max_lengths = [0] * len(table.columns)
    
    # Iterate over each row in the table
    for row in table.rows:
        # Iterate over each cell in the row
        for col_index, cell in enumerate(row.cells):
            # Get the text content of the cell
            cell_text = cell.text.strip()
            
            # Update the maximum length for the column if needed
            if len(cell_text) > max_lengths[col_index]:
                max_lengths[col_index] = len(cell_text)
            
    return max_lengths


def search_text_in_document(input_docx_path, search_text):
    # Load the document from the specified path
    doc = Document(input_docx_path)

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Check if the search text is present in the paragraph text
        if search_text in paragraph.text:
            # If found, return True
            return True

    # If the search text is not found in any paragraph, return False
    return False


def change_page_size(input_docx_path, output_docx_path, page_width, page_height):
    # Load the input DOCX document
    doc = Document(input_docx_path)
    
    # Convert page width and height from inches to Twips (twentieths of a point)
    page_width_twips = Inches(page_width)
    page_height_twips = Inches(page_height)

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Get the current line spacing of the paragraph
        current_spacing = paragraph.paragraph_format.line_spacing
        print('ttttttttttttttttttttt current spacing is : ', current_spacing)

        if current_spacing:
            paragraph.paragraph_format.line_spacing = Pt(1)
    
    # Iterate through all sections in the document
    for section in doc.sections:
        # Set the page width and height for the section
        # section.page_width = page_width_twips
        # section.page_height = page_height_twips
        section.page_width = section.page_width 
        section.page_height = int(section.page_height * 1)
    
    # Save the modified document to the output path
    doc.save(output_docx_path)


def adjust_margins(docx_path):
    doc = Document(docx_path)

    # Get the sections in the document
    sections = doc.sections

    # Iterate through each section
    for section in sections:
        

        # Get the margins of the section
        left_margin = section.left_margin
        right_margin = section.right_margin
        top_margin = section.top_margin
        bottom_margin = section.bottom_margin

        # Print the margins of the section
        # print("YYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYYY Margins for Section:")
        print(f"Left: {left_margin}, Right: {right_margin}, Top: {top_margin}, Bottom: {bottom_margin}")
        section.left_margin //= 2
        section.right_margin //= 2
        section.top_margin //= 2
        section.bottom_margin //= 2

        # Get the margins of the section
        left_margin = section.left_margin
        right_margin = section.right_margin
        top_margin = section.top_margin
        bottom_margin = section.bottom_margin

        # Print the margins of the section
        # print("ZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZZ Margins for Section:")
        print(f"Left: {left_margin}, Right: {right_margin}, Top: {top_margin}, Bottom: {bottom_margin}")

    doc.save(docx_path)


def reduce_font_size(docx_path):
    doc = Document(docx_path)

    # Define the reduction factor for font size (50% reduction)
    reduction_factor = 0.95

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Iterate through each run in the paragraph
        for run in paragraph.runs:
            # Get the current font size
            current_size = run.font.size

            # Reduce the font size by the reduction factor
            if current_size:
                new_size = Pt(current_size.pt * reduction_factor)

                # Set the new font size
                run.font.size = new_size

    # Save the modified document
    doc.save(docx_path)
